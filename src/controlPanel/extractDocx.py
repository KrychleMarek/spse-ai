import re
import os
from pathlib import Path
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
import json
from openai import OpenAI

ROOT_DIR = Path(__file__).resolve().parent.parent.parent
INPUT_DIR = ROOT_DIR / "data" / "raw"
OUTPUT_DIR = ROOT_DIR / "data" / "processed" / "embeddFiles"
CONFIG_PATH = ROOT_DIR / "config.json"

with open(CONFIG_PATH) as f:
    config = json.load(f)

client = OpenAI(api_key=config["api_key"]) 



def iter_block_items(parent): # Průchod dokumentem po blocích či odstavcích či tabulkek čičičičiči meow
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Neznámý rodičovský element")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def get_subject_map_from_table(doc):
    # Zkouší najít "Výčet vyučovacíh předmětů" ze které sežene názvy a zkratky předmětů a vyprdne je jako mapu
    subject_map = {}
    seen_subjects = set()
    for table in doc.tables:
        is_target_table = False
        name_col_index = -1
        abbr_col_index = -1
        
        for row_idx, row in enumerate(table.rows[:5]):
            for col_idx, cell in enumerate(row.cells):
                txt = cell.text.lower()
                if "název" in txt and "předmětu" in txt:
                    is_target_table = True
                    name_col_index = col_idx
                if "zkr" in txt or "zkratka" in txt:
                    abbr_col_index = col_idx
            if is_target_table: break
        
        if is_target_table:
            for row in table.rows[row_idx + 1:]:
                if len(row.cells) > name_col_index:
                    raw_name = clean_text(row.cells[name_col_index].text)
                    raw_abbr = ""
                    if abbr_col_index != -1 and len(row.cells) > abbr_col_index:
                        raw_abbr = clean_text(row.cells[abbr_col_index].text)

                    if raw_name and "Celkem" not in raw_name:
                        if raw_name not in seen_subjects:
                            subject_map[raw_name] = raw_abbr
                            seen_subjects.add(raw_name)
            return subject_map # Mapa objektů př. {"Fyzika" : "Fyz"}
    return {}



def simplify_text_with_gpt(text_list, subject_name):
    global client
    
    if not client:
        return "\n".join(text_list)

    full_text = "\n\n".join(text_list)
    
    # Prostě prompt no. Pokuď se neplteu tak docela žere peníze. Bude se to teda pouštět jenom jednou za rok ale jenom za jeden obor to bylo 0.05$ takže to je cca 0.3$ za všechny takže vlastně nic. Bohužel se neůžu ani na dashboard podívat kolik mi zbývá. Je 10.12.2025 zakže typuju že tam bude tak cca 4.85$ ¯\_(ツ)_/¯ ale tak jsou to skolni prachy tak mi to muze byt jedno
    prompt = f"""
    Zestručni text, který popisuje pojetí vyučovacího předmětu '{subject_name}'.
    
    Úkoly:
    1. Přepiš text do srozumitelné formy a zachovej klíčové body.
    2. Formátuj výstup pomocí odstavců a odrážek.
    3. Používej POUZE informace z původního textu. NEVYMÝŠLEJ SI.

    STRIKTNÍ PRAVIDLA PRO FORMÁTOVÁNÍ:
    - Nadpisy piš jako čistý text na samostatný řádek.
    - U nadpisů NEPOUŽÍVEJ mřížky (#), hvězdičky (*), odrážky (•) ani žádné jiné formátovací znaky.
    - Nadpisy nesmí být tučné ani podtržené.
    - Ukonči odpověď hned po posledním bodu. Nepřidávej žádné oddělovací čáry, podtržítka ani jiné dekorativní znaky na konec textu

    Struktura výstupu bude vypadat přesně takto:
    Co se naučím v {subject_name}
    [Zde bude text cíle]

    Co se vyučuje:
    [Obsah]

    Co se naučíš:
    [Obsah]

    Jak probíhá výuka:
    [Obsah]

    Proč je předmět důležitý:
    [Obsah]

    

    
    [[ZAČÁTEK PŮVODNÍHO TEXTU]]
    {full_text}
    [[KONEC PŮVODNÍHO TEXTU]]]
    """
    
    print(f" -> Volám GPT pro zjednodušení textu pro předmět: {subject_name}...")
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Jsi zkušený editor vzdělávacích textů. Tvojí úlohou je zjednodušovat složité kurikulární dokumenty."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        
        simplified_text = response.choices[0].message.content.strip()
        print(" -> GPT zjednodušení hotovo.")
        simplified_text = simplified_text.replace("#", "") # Pojištění že tam nebudou ty pitomý hastagy
        return simplified_text
        
    except Exception as e:
        print(f"Chyba při volání GPT API pro {subject_name}: {e}. Používám původní text.")
        return full_text


def parse_document(docx_path):
    doc = Document(docx_path)
    subject_map_overview = get_subject_map_from_table(doc)
    
    subjects_details = []
    current_subject = None
    current_year = None
    capturing_concept = False
    
    re_subj_start = re.compile(r"Název vyučovacího předmětu:\s*(.+)", re.IGNORECASE)
    re_toc_entry = re.compile(r'\d+$')
    re_concept_start = re.compile(r"Pojetí vyučovacího předmětu", re.IGNORECASE)
    re_curriculum_start = re.compile(r"Rozpis učiva a výsledků vzdělávání", re.IGNORECASE)
    re_year = re.compile(r"^(\d+)\.\s*ročník", re.IGNORECASE)
    re_page_number = re.compile(r'^\d+$')
    re_abbr_in_title = re.compile(r'\((.*?)\)')

    print("Zahajuji extrakci obsahu...")

    for block in iter_block_items(doc): 
        if isinstance(block, Paragraph): # Pokuď to je normální text
            text = clean_text(block.text)
            if not text: continue

            match_subj = re_subj_start.search(text)
            if match_subj:
                raw_heading = match_subj.group(1).strip()
                if re_toc_entry.search(raw_heading): continue

                if current_subject:
                    subjects_details.append(current_subject)
                
                name = re.sub(r'\s*\(.*?\)', '', raw_heading).strip()
                abbr_match = re_abbr_in_title.search(raw_heading)
                abbr = abbr_match.group(1) if abbr_match else ""
                if not abbr and name in subject_map_overview:
                    abbr = subject_map_overview[name]

                current_subject = { #Jeden předmět se všim
                    "name": name,
                    "abbr": abbr,
                    "concept_text": [],
                    "curriculum": {1: [], 2: [], 3: [], 4: []}
                }
                current_year = None
                capturing_concept = False
                continue

            if not current_subject: continue

            if re_concept_start.search(text): # Prohledává obsah odstavce pro nadpis po kterym začne extrakci
                capturing_concept = True
                continue 
            if re_curriculum_start.search(text): # Opak původního ifu
                capturing_concept = False
                continue

            if capturing_concept:
                if not re_page_number.match(text):
                    current_subject["concept_text"].append(text)
                continue 

            match_year = re_year.search(text)
            if match_year:
                current_year = int(match_year.group(1))
                continue

        elif isinstance(block, Table): # Pokudď je tabulka
            if current_subject and current_year and not capturing_concept:
                ucivo_idx = -1
                if not block.rows: continue
                
                header_row = block.rows[0]
                for i, cell in enumerate(header_row.cells):
                    if "učivo" in cell.text.lower():
                        ucivo_idx = i
                        break
                
                if ucivo_idx != -1:
                    for row in block.rows[1:]:
                        if len(row.cells) > ucivo_idx:
                            cell_text = clean_text(row.cells[ucivo_idx].text)
                            items = [x.strip() for x in re.split(r'[;•\n]', cell_text) if len(x.strip()) > 1]
                            current_subject["curriculum"][current_year].extend(items)

    if current_subject:
        subjects_details.append(current_subject)

    return subjects_details # Všechny předměty po extrakci


def write_combined_docx(subjects_details, filepath_name, filename):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading(filename, 0)
    
    # Napdis seznam předmětu pro chunker. MUSÍ SE VŽDY SCHODOVAT
    doc.add_heading('JAKÉ PŘEDMĚTY JSOU/SE VYUČUJÍ NA OBORU:', 1)
    
    for det in subjects_details:
        display_name = f"{det['name']} ({det['abbr']})" if det['abbr'] else det['name']
        doc.add_paragraph(display_name, style='List Bullet')
    
    doc.add_page_break()

    # Výpis každýho předmětu
    for detail in subjects_details:
        full_title = f"{detail['name']} ({detail['abbr']})" if detail['abbr'] else detail['name']
        
        # Nadpis předmětu
        doc.add_heading(full_title, 1)
        
        concept_lines = detail.get("concept_text", [])
        if concept_lines:
            simplified_text = simplify_text_with_gpt(concept_lines, full_title)
            
            simplified_paragraphs = [p.strip() for p in simplified_text.split('\n\n') if p.strip()]

            for paragraph in simplified_paragraphs:
                if re.match(r'^\s*[*•-]\s*', paragraph):
                    list_items = re.split(r'[*•-]', paragraph)
                    for item in list_items:
                        item = item.strip()
                        if item:
                            doc.add_paragraph(item, style='List Bullet')
                # Zpracování potenciálních podnadpisů
                elif len(paragraph) < 60 and not paragraph.endswith('.'):
                    p = doc.add_paragraph(paragraph)
                    p.style = 'Heading 3'
                else:
                    doc.add_paragraph(paragraph)

            doc.add_paragraph() # Mezera

        # Pokud má ročník tak true protože z nějakýho pytomího důvodu "existuje" předmět "Rozvoj Osobnosti (ROZ)"
        # kterej nikdo na celý škole podle mě nemá a nebude mít. Certified spše classic
        has_curriculum = any(detail["curriculum"].get(year) for year in range(1, 5)) 
        
        if has_curriculum:
            for year in range(1, 5):
                items = detail["curriculum"].get(year, [])
                items = list(dict.fromkeys(items))

                if items:
                    # Napdis pro chunker: "Co se učí X. ročník v Název Předmětu (ZKR)"
                    doc.add_heading(f"Co se učí {year}. ročník v {full_title}", 2)
                    for item in items:
                        doc.add_paragraph(item, style='List Bullet')
        
        doc.add_paragraph()
        doc.add_page_break()

    doc.save(filepath_name)
    print(f"HOTOVO. Soubor uložen jako: {filepath_name}")


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    docx_files = list(INPUT_DIR.glob("*.docx"))
    
    if not docx_files:
        print(f"Chyba: V adresáři {INPUT_DIR} nebyly nalezeny žádné soubory .docx.")
    else:
        print(f"Nalezeno {len(docx_files)} souborů .docx ke zpracování.")

        for INPUT_DIR_path in docx_files:
            output_filename = f"{INPUT_DIR_path.stem}.docx"
            output_file_path = OUTPUT_DIR / output_filename

            print("-" * 50)
            print(f"Zpracovávám soubor: {INPUT_DIR_path.name}")
            
            try:
                details = parse_document(INPUT_DIR_path) 
                write_combined_docx(details, output_file_path, output_filename)
            
            except Exception as e:
                print(f"CHYBA při zpracování souboru {INPUT_DIR_path.name}: {e}")